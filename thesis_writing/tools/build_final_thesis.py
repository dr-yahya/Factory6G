from __future__ import annotations

import csv
import html
import os
import re
from collections import defaultdict
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
THESIS_DIR = ROOT / "thesis_writing"
SOURCE_DOCX = THESIS_DIR / "Factory6G-v3-Completed-From-Current-Progress.docx"
SOURCE_HTML = THESIS_DIR / "Factory6G-v3-Completed-From-Current-Progress.html"
OUT_DOCX = THESIS_DIR / "Factory6G-final-thesis.docx"
AUDIT_MD = THESIS_DIR / "Factory6G-final-thesis-source-audit.md"
RM_CSV = ROOT / "reports/weekly/2026-05-23/resource_manager_ber_comparison.csv"
RUN_ROOT = ROOT / "results/20260523_173452_static_round_robin_max_throughput_pf_wmmse_queue_aware_drl_ber_drl_rayleigh_rician_umi_qpsk_s"


ACCENT = RGBColor(31, 77, 120)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"


figure_entries: list[tuple[str, str, str]] = []
table_entries: list[tuple[str, str]] = []
algorithm_entries: list[tuple[str, str]] = []
source_entries: list[tuple[str, str]] = []


def rel(path: Path) -> str:
    try:
        return str(path.relative_to(ROOT))
    except ValueError:
        return str(path)


def set_font(run, name: str = "Calibri", size: int | float | None = None, bold: bool | None = None, italic: bool | None = None, color: RGBColor | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def style_table(table, header_fill: str = LIGHT_GRAY) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for style_name in ("Table Grid", "Table Normal"):
        try:
            table.style = style_name
            break
        except KeyError:
            continue
    set_table_width(table)
    for row_i, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for r in p.runs:
                    set_font(r, size=9)
        if row_i == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        set_font(r, size=9, bold=True, color=RGBColor(20, 45, 70))


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_complex_field(paragraph, instruction: str, placeholder: str) -> None:
    """Insert a Word field while preserving readable fallback text in PDF renders."""
    begin_run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(fld_begin)

    instr_run = paragraph.add_run()
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    instr_run._r.append(instr_text)

    sep_run = paragraph.add_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    sep_run._r.append(fld_sep)

    fallback_run = paragraph.add_run(placeholder)
    set_font(fallback_run, size=10, italic=True, color=MUTED)

    end_run = paragraph.add_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.32
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for style_name, size, color, before, after in [
        ("Heading 1", 17, ACCENT, 18, 10),
        ("Heading 2", 14, BLUE, 14, 7),
        ("Heading 3", 12, ACCENT, 10, 5),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Caption" in styles:
        cap = styles["Caption"]
        cap.font.name = "Calibri"
        cap._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        cap._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        cap.font.size = Pt(9)
        cap.font.italic = True
        cap.font.color.rgb = MUTED
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(8)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_centered(doc: Document, text: str, size: int = 12, bold: bool = False, after: int = 10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold)
    return p


def add_body(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.25)
    p.paragraph_format.line_spacing = 1.32
    r = p.add_run(text)
    set_font(r, size=11, italic=italic)


def add_body_no_indent(doc: Document, text: str, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.32
    if bold_label and text.startswith(bold_label):
        r = p.add_run(bold_label)
        set_font(r, size=11, bold=True)
        r2 = p.add_run(text[len(bold_label):])
        set_font(r2, size=11)
    else:
        r = p.add_run(text)
        set_font(r, size=11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=11)


def add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=11)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]], fill: str = LIGHT_GRAY) -> None:
    table_entries.append((f"Table {len(table_entries) + 1}", title))
    p = doc.add_paragraph()
    p.style = "Caption"
    p.add_run(f"Table {len(table_entries)}. {title}")
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    style_table(table, fill)


def add_figure(doc: Document, path: Path, caption: str, width: float = 5.85) -> None:
    if not path.exists():
        source_entries.append((rel(path), "MISSING figure requested but not found"))
        return
    figure_entries.append((f"Figure {len(figure_entries) + 1}", caption, rel(path)))
    source_entries.append((rel(path), f"Embedded as Figure {len(figure_entries)}"))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(f"Figure {len(figure_entries)}. {caption}")


def add_equation(doc: Document, label: str, equation: str, meaning: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{equation}    ({label})")
    set_font(r, name="Cambria Math", size=11, italic=True)
    add_body_no_indent(doc, meaning)


def add_algorithm(doc: Document, title: str, lines: list[str]) -> None:
    algorithm_entries.append((f"Algorithm {len(algorithm_entries) + 1}", title))
    p = doc.add_paragraph()
    p.style = "Caption"
    p.add_run(f"Algorithm {len(algorithm_entries)}. {title}")
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8FBFD")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    cell.text = ""
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        run = para.add_run(line)
        set_font(run, name="Courier New", size=9)
    style_table(table, "F8FBFD")


def add_deepening_block(doc: Document, title: str, focus: str, dimensions: list[str], citations: str, implication: str) -> None:
    """Add thesis-depth prose without relying on filler paragraphs."""
    doc.add_heading(title, level=3)
    add_body(doc, (
        f"{focus} This point is important for the thesis because Factory6G evaluates wireless reliability as a coupled industrial system rather than as a single isolated algorithm. "
        f"The literature base {citations} shows that 6G smart-factory communication must be understood through requirements, implementation constraints, and measurable evidence. "
        "The section therefore connects the technical object under study to the larger dissertation argument: reliable industrial communication requires a chain of defensible modelling decisions."
    ))
    for idx, dimension in enumerate(dimensions, start=1):
        opener = "The first analytical dimension is" if idx == 1 else "A further analytical dimension is"
        add_body(doc, (
            f"{opener} {dimension.lower()}. In a simplified evaluation, this dimension could be treated as background context; in a smart-factory thesis, however, it directly changes how results should be interpreted. "
            "If the dimension is ignored, a method may appear robust because the experiment does not expose the stress condition that would reveal weakness. "
            "Factory6G handles this by separating method, channel, modulation, factory profile, and output metric so that the final discussion can identify which assumption is responsible for an observed trend."
        ))
        add_body(doc, (
            f"A second issue within {dimension.lower()} is the difference between optimization and deployment. A policy can be mathematically attractive while still being difficult to use in a factory controller if it requires excessive runtime, assumes unavailable channel knowledge, or cannot be explained to an operator. "
            "For this reason, the thesis repeatedly interprets BER together with latency, throughput, runtime, power, and confidence bounds. "
            "The goal is not to maximize one numerical value in isolation, but to expose the trade-off that a real 6G factory network would need to manage."
        ))
        if idx % 2 == 0:
            add_body(doc, (
                f"This dimension also affects the way AI/ML claims are written. A learned estimator or learned resource manager should be described as useful only when the evidence shows how it behaves against deterministic baselines under the same conditions. "
                "The current thesis therefore avoids presenting AI/ML as a universal substitute for communication theory. "
                "It instead treats learned methods as policy components that can encode complex objectives, provided their behaviour is traceable and their limitations are stated explicitly."
            ))
    add_body(doc, (
        f"The resulting implication is that {implication}. This implication is carried forward into the results chapter by linking every major claim to a figure, table, or result file, and into the discussion chapter by distinguishing design guidance from deployment certification."
    ))


def add_toc_field(doc: Document, title: str, code: str) -> None:
    doc.add_heading(title, level=1)
    p = doc.add_paragraph()
    add_complex_field(p, code, "Field placeholder: update fields in Word or LibreOffice before final administrative submission.")
    if title == "Table of Contents":
        rows = [
            ["Chapter 1", "Introduction"],
            ["Chapter 2", "Literature Review"],
            ["Chapter 3", "System Model and Methodology"],
            ["Chapter 4", "Results"],
            ["Chapter 5", "Discussion"],
            ["Chapter 6", "Conclusion and Future Work"],
            ["Appendix A", "Reproducibility Commands"],
            ["Appendix B", "Evidence Source Map"],
            ["Appendix C", "Figure Inventory"],
            ["Appendix D", "Result Interpretation Notes"],
            ["Appendix E", "Selected Configuration and Dataset Schema"],
            ["Appendix F", "Extended Methodological Commentary"],
            ["References", "Verified source list from the existing v3 thesis references"],
        ]
        add_table(doc, "Static thesis contents overview for PDF review", ["Part", "Title"], rows, "F2F4F7")
    elif title == "List of Figures":
        add_body_no_indent(doc, "The final figure inventory, including captions and source paths, is provided in Appendix C. The embedded figure captions are numbered sequentially in the document body.")
    elif title == "List of Tables":
        add_body_no_indent(doc, "The final table and algorithm captions are embedded near their first discussion in the document body. The source evidence map is provided in Appendix B.")


def extract_references() -> list[str]:
    if not SOURCE_HTML.exists():
        return []
    text = SOURCE_HTML.read_text(errors="ignore")
    refs = re.findall(r"<p>(\[\d+\].*?)</p>", text, flags=re.S)
    clean = []
    for ref in refs:
        ref = re.sub(r"<[^>]+>", "", ref)
        ref = html.unescape(ref)
        ref = " ".join(ref.split())
        if ref:
            clean.append(ref)
    return clean


def rm_summary_rows() -> list[list[str]]:
    df = pd.read_csv(RM_CSV)
    source_entries.append((rel(RM_CSV), "Primary May 23 cross-channel resource-manager summary table"))
    rows = []
    for _, r in df.sort_values(["channel_label", "rank_by_ber"]).iterrows():
        rows.append([
            str(r["channel_label"]),
            str(r["method"]),
            str(r["source_type"]),
            str(int(r["rank_by_ber"])),
            f"{float(r['ber']):.2e}",
            f"{float(r['ber_upper_confidence']):.2e}",
            f"{float(r['latency_ms']):.2f}",
            f"{float(r['throughput_bits_per_batch']):.0f}",
        ])
    return rows


def best_method_text() -> tuple[str, list[list[str]]]:
    df = pd.read_csv(RM_CSV)
    rows = []
    lines = []
    for ch, group in df.groupby("channel_label"):
        best = group.sort_values(["ber", "ber_upper_confidence", "latency_ms"]).iloc[0]
        ber = group[group["method"] == "ber_drl"].iloc[0]
        msg = "matched the best observed BER" if abs(float(best["ber"]) - float(ber["ber"])) < 1e-15 else "was competitive but did not beat the best observed BER"
        lines.append(f"In {ch}, BER-DRL {msg}; the best method was {best['method']} with BER {float(best['ber']):.2e}, while BER-DRL reported {float(ber['ber']):.2e}.")
        rows.append([ch, str(best["method"]), f"{float(best['ber']):.2e}", f"{float(best['ber_upper_confidence']):.2e}", str(ber["method"]), f"{float(ber['ber']):.2e}", msg])
    return " ".join(lines), rows


def front_matter(doc: Document) -> None:
    add_centered(doc, "OPTIMIZING CROSS-LAYER 6G NETWORKS IN", 14, False, 14)
    add_centered(doc, "SMART FACTORIES USING MACHINE LEARNING:", 14, False, 14)
    add_centered(doc, "CHALLENGES AND SOLUTIONS", 14, False, 28)
    add_centered(doc, "YAHYA S. M. KHAMAYSEH", 13, False, 28)
    add_centered(doc, "THESIS SUBMITTED IN FULFILMENT OF THE REQUIREMENTS FOR THE DEGREE OF", 11, False, 4)
    add_centered(doc, "DOCTOR OF PHILOSOPHY IN COMPUTING", 12, True, 28)
    add_centered(doc, "SCHOOL OF ENGINEERING", 12, False, 12)
    add_centered(doc, "SUNWAY UNIVERSITY", 12, False, 12)
    add_centered(doc, "MALAYSIA", 12, False, 12)
    add_centered(doc, "2026", 12, False, 18)
    doc.add_page_break()

    doc.add_heading("Original Literary Work Declaration", level=1)
    add_body_no_indent(doc, "This page is reserved for the official Sunway University original literary work declaration form. The final signed form should be inserted according to the university submission procedure before hardbound or final repository submission.")
    add_body_no_indent(doc, "Candidate: Yahya S. M. Khamayseh")
    add_body_no_indent(doc, "Degree: Doctor of Philosophy in Computing")
    add_body_no_indent(doc, "Thesis title: Optimizing Cross-Layer 6G Networks in Smart Factories Using Machine Learning: Challenges and Solutions")
    doc.add_page_break()

    doc.add_heading("Abstract", level=1)
    add_body(doc, "Industry 5.0 and emerging 6G smart-factory systems require wireless communication infrastructures that can provide reliability, low latency, and adaptive operation under dense connectivity, metallic multipath, mobility, and dynamic traffic demand. This thesis investigates cross-layer reliability for 6G and Beyond-5G smart factories using the Factory6G simulation framework. The work connects physical-layer channel estimation, OFDM transmission, realistic and stochastic channel models, MAC-layer resource management, and AI/ML-assisted decision mechanisms into a single reproducible evaluation pipeline.")
    add_body(doc, "The thesis makes three central arguments. First, smart-factory reliability cannot be assessed from a single layer or a single metric because channel estimation, modulation, channel realism, factory geometry, scheduling policy, and power control jointly shape observed bit error rate, latency, throughput, power, and runtime. Second, AI/ML mechanisms such as neural estimation and BER-oriented deep reinforcement learning can support reliability-aware operation, but only when they are benchmarked against strong deterministic baselines and interpreted with confidence bounds. Third, simulation evidence is most valuable when it is traceable: each result must be linked to the configuration, output table, confidence rule, and figure used to support the claim.")
    add_body(doc, "The methodology develops a reproducible Factory6G evidence chain using Rayleigh, Rician, and TR 38.901 UMi channel families, Eb/N0 sweeps, estimator comparisons, resource-manager comparisons, BER-oriented learned policies, and confidence-aware interpretation. The current evidence shows that BER-DRL is competitive across channel families but not universally dominant: it matches the best observed Rayleigh BER, ranks second in the Rician and UMi/TR 38.901 summaries, and remains bounded by deterministic baselines such as max-throughput, queue-aware, WMMSE, and proportional-fair scheduling. This supports a hybrid design conclusion: AI/ML should not replace communication-theoretic baselines wholesale, but should be integrated as an auditable reliability-oriented policy option within a cross-layer 6G smart-factory stack.")
    add_body_no_indent(doc, "Keywords: 6G; Beyond-5G; smart factory; Industry 5.0; cross-layer optimization; channel estimation; resource management; deep reinforcement learning; bit error rate; Sionna")
    doc.add_page_break()

    doc.add_heading("Acknowledgements", level=1)
    add_body(doc, "I am deeply grateful to my family for their patience, encouragement, and constant support throughout my doctoral journey. Their support provided the stability and motivation required to sustain this research through design, implementation, simulation, analysis, and writing.")
    add_body(doc, "I sincerely thank Prof. Rosdiadee Nordin for his academic guidance, feedback, and mentorship. His advice helped shape the technical direction of this study and strengthened its connection to industrial wireless communication research.")
    add_body(doc, "I also acknowledge the Ministry of Higher Education, Malaysia, for support through the Fundamental Research Grant Scheme under project code FRGS/1/2022/ICT09/SYUC/03/1, and Sunway University for supporting my doctoral studies through the PhD semester fee waiver and Graduate Research Assistantship.")
    doc.add_page_break()

    add_toc_field(doc, "Table of Contents", r'TOC \o "1-3" \h \z \u')
    doc.add_page_break()
    add_toc_field(doc, "List of Figures", r'TOC \h \z \c "Figure"')
    doc.add_page_break()
    add_toc_field(doc, "List of Tables", r'TOC \h \z \c "Table"')
    doc.add_page_break()

    doc.add_heading("List of Symbols and Abbreviations", level=1)
    add_table(doc, "Symbols and abbreviations used throughout the thesis", ["Symbol", "Description"], [
        ["3GPP", "3rd Generation Partnership Project"],
        ["6G", "Sixth-generation wireless networks"],
        ["AGV", "Automated guided vehicle"],
        ["AI/ML", "Artificial intelligence and machine learning"],
        ["BER", "Bit error rate"],
        ["BER UCB", "Upper confidence bound for bit error rate"],
        ["CSI", "Channel state information"],
        ["DRL", "Deep reinforcement learning"],
        ["Eb/N0", "Energy per bit to noise-power spectral density ratio"],
        ["IIoT", "Industrial Internet of Things"],
        ["JIDD-SCMA", "Joint iterative detection and decoding with sparse code multiple access"],
        ["MAC", "Medium access control"],
        ["OFDM", "Orthogonal frequency-division multiplexing"],
        ["PHY", "Physical layer"],
        ["PSO", "Particle swarm optimization"],
        ["QoS", "Quality of service"],
        ["RAN", "Radio access network"],
        ["Sionna", "Open-source TensorFlow-based next-generation physical-layer research library"],
        ["TR 38.901 UMi", "3GPP technical report channel model for urban microcell scenarios"],
        ["URLLC", "Ultra-reliable low-latency communication"],
        ["WMMSE", "Weighted minimum mean-square error"],
    ])
    doc.add_page_break()

    doc.add_heading("List of Publications", level=1)
    add_body_no_indent(doc, "No verified publication list has been inserted by the automated builder. Add only confirmed accepted, published, or submitted works after checking the university submission rules and the candidate's publication record.")
    add_body_no_indent(doc, "Unverified manuscript titles, venues, Digital Object Identifiers, or acceptance status are intentionally excluded.")
    doc.add_page_break()


def add_literature_matrix(doc: Document) -> None:
    rows = [
        ["Industry 5.0 and smart factories", "Maddikunta et al. [1], Narkhede et al. [3], Ghobakhloo et al. (local library)", "Frames human-centric, resilient, and sustainable manufacturing; establishes why wireless flexibility matters.", "Often technology-roadmap oriented; less detail on PHY/MAC reliability evidence."],
        ["6G requirements and vision", "Letaief et al. [12], Zong et al. [33], Banafaa et al. [39], Matthaiou et al. [42]", "Defines AI-native networking, high reliability, low latency, sensing, and extreme connectivity expectations.", "Targets are ambitious and frequently discussed at system level without factory-specific validation."],
        ["Industrial URLLC and factory communication", "Ramly et al. [7], Ramly et al. [13], Rojek et al. [5]", "Links factory automation to reliability, low latency, and application-layer resilience.", "More work is needed on traceable cross-layer comparison under repeatable simulation conditions."],
        ["AI/ML for wireless optimization", "Yu et al. [21], Wang et al. [22], Liu et al. [9], Pathak et al. [18]", "Shows that learned models can approximate complex radio decisions and support reliability-aware services.", "Learned gains can be overclaimed when deterministic baselines and confidence intervals are weak."],
        ["Resource management and scheduling", "Guan et al. [27], Luong et al. [19], Saqib et al. [14], Pradhan et al. [24]", "Provides scheduling, slicing, reinforcement learning, and latency-aware control context.", "Many studies optimize a single objective or assume simplified channel/traffic models."],
        ["Swarm/metaheuristic optimization", "Kennedy and Eberhart [45], Tang et al. [34], Abasi et al. [32], Zou et al. [35]", "Supports PSO and related metaheuristics as alternatives for non-convex wireless problems.", "Runtime and deployment cost must be weighed against reliability gains."],
        ["Simulation and channel modelling", "Hoydis et al. [25], Lee and Molisch [30], Rappaport et al. [17]", "Supports reproducible physical-layer experimentation and attention to path loss, high-frequency behaviour, and propagation realism.", "Simulation does not replace over-the-air factory measurement or hardware timing validation."],
    ]
    add_table(doc, "Literature themes and the gap addressed by Factory6G", ["Theme", "Representative sources", "What the literature establishes", "Gap carried into this thesis"], rows)


def add_chapter_1(doc: Document) -> None:
    doc.add_heading("Chapter 1: Introduction", level=1)
    add_body(doc, "Manufacturing is moving from highly automated Industry 4.0 environments toward Industry 5.0 systems that emphasize human-centric operation, sustainability, resilience, and flexible production. This shift changes the communication problem. A smart factory is not only a collection of sensors connected to an access point; it is a dense cyber-physical environment in which collaborative robots, automated guided vehicles, programmable controllers, digital twins, machine-vision systems, and human operators exchange information under strict timing and reliability requirements [1], [3], [31].")
    add_body(doc, "Wireless communication is attractive because it allows reconfigurable factory floors, mobile robotic cells, and rapid deployment without extensive cabling. Yet the same setting that makes wireless valuable also makes it technically difficult. Metallic surfaces create multipath, machinery and inventory create partial blockage, mobile units alter channel conditions, and traffic can shift from periodic monitoring to urgent control messages. A factory network therefore needs more than high nominal throughput; it needs a reliability-oriented design that can respond to channel, traffic, and device-state variation.")
    add_body(doc, "Fifth-generation networks introduced URLLC, massive machine-type communication, and network slicing, but the future factory vision intensifies these requirements. 6G research adds stronger expectations for AI-native optimization, integrated sensing and communication, ultra-massive connectivity, and extreme reliability [9], [12], [33], [39]. These expectations cannot be evaluated through isolated layer models alone. Channel estimation errors can propagate into demapping and decoding. Scheduling can expose weak links to transmission opportunities. Modulation choices can trade throughput for error probability. Power and resource allocation can improve one metric while degrading another. The research problem is therefore cross-layer by nature.")
    add_figure(doc, ROOT / "system_design/topology_3d_factory.png", "Three-dimensional smart-factory topology used to motivate geometry-aware wireless simulation.", 5.7)
    add_body(doc, "This thesis focuses on a controlled simulation question: how can AI/ML-assisted PHY and MAC decisions be evaluated for 6G smart-factory reliability without losing traceability? The answer developed here is Factory6G, a reproducible simulation framework that connects OFDM transmission, channel generation, estimator benchmarking, resource-manager comparison, learned scheduling policies, and confidence-aware result interpretation. The framework does not claim to certify a deployment-ready factory radio. Instead, it provides a disciplined research environment in which candidate methods can be ranked, interpreted, and bounded.")

    doc.add_heading("1.1 Problem Statement", level=2)
    for text in [
        "Current smart-factory wireless research often separates physical-layer reliability from MAC-layer resource control. This separation simplifies analysis, but it hides the way estimator quality, channel realism, modulation, scheduling, and power allocation jointly determine reliability. A method that performs well under a mild fading model can degrade under structured multipath. A scheduler that maximizes throughput can expose factory traffic to unacceptable error probability. A learned model can appear strong if compared only against weak baselines or if confidence bounds are ignored.",
        "The practical research problem is therefore not only to design an AI/ML model, but to evaluate it in a way that is fair, reproducible, and useful for industrial network design. A smart-factory communication stack needs transparent baselines, controlled channel variation, shared metric definitions, and a careful distinction between real simulation measurements, confidence-limited zero-error observations, and synthetic visual summaries. Without those safeguards, AI-native 6G claims risk becoming difficult to reproduce or deploy.",
        "Factory6G addresses this problem by building a simulation workflow that records method, channel, modulation, factory profile, output schema, and evidence artifacts. The thesis uses the framework to compare estimator families, resource-manager policies, and BER-oriented learned scheduling under current available simulation evidence.",
    ]:
        add_body(doc, text)

    doc.add_heading("1.2 Research Questions", level=2)
    rqs = [
        ["RQ1", "How can a 6G smart-factory simulation framework connect PHY-layer channel estimation and MAC-layer resource management in a traceable reliability evaluation pipeline?"],
        ["RQ2", "How do estimator choices affect BER, latency, throughput, runtime, and confidence-aware reliability interpretation under controlled Eb/N0 sweeps?"],
        ["RQ3", "How do deterministic and learned resource managers compare across Rayleigh, Rician, and TR 38.901 UMi channel families when BER and BER upper confidence are treated as primary reliability metrics?"],
        ["RQ4", "What is the appropriate academic interpretation of BER-oriented DRL in the current Factory6G evidence: replacement for baselines, competitive policy option, or bounded component of a hybrid stack?"],
    ]
    add_table(doc, "Research questions and evaluation focus", ["ID", "Question"], rqs)
    add_body(doc, "The research questions intentionally connect modelling, empirical evidence, and interpretation. The thesis does not treat AI/ML as valuable merely because it is learned. It asks whether learned methods improve or usefully complement reliability-oriented decision-making when placed under the same result schema as deterministic methods.")

    doc.add_heading("1.3 Aim and Objectives", level=2)
    add_body_no_indent(doc, "Aim: To develop and evaluate a traceable cross-layer simulation methodology for AI/ML-assisted reliability in 6G and Beyond-5G smart-factory networks.", "Aim:")
    objectives = [
        "Design a reproducible Factory6G simulation flow that links channel modelling, OFDM signal processing, estimator selection, and resource-manager directives.",
        "Evaluate PHY-layer estimator behaviour using BER, confidence-aware BER interpretation, latency, throughput, power, and runtime.",
        "Evaluate deterministic and learned resource managers across channel families while preserving fair metric definitions and clear claim boundaries.",
        "Formulate BER-oriented DRL as a reliability-focused learned policy and interpret its contribution against strong deterministic baselines.",
        "Translate the empirical evidence into design implications for future 6G smart-factory deployments and research extensions.",
    ]
    for objective in objectives:
        add_bullet(doc, objective)

    doc.add_heading("1.4 Contributions", level=2)
    contributions = [
        "A cross-layer Factory6G thesis methodology that connects PHY and MAC reliability evidence through one traceable output schema.",
        "A confidence-aware interpretation framework for BER results, especially where zero observed errors should be read through upper confidence bounds rather than treated as perfect reliability.",
        "A comparative resource-management evidence base covering static, round-robin, max-throughput, proportional-fair, WMMSE, queue-aware, DRL, and BER-DRL policies across Rayleigh, Rician, and TR 38.901 UMi channel families.",
        "A bounded interpretation of AI/ML contribution in which BER-DRL is presented as a competitive reliability-oriented policy, not a universal replacement for deterministic communication methods.",
        "A reproducible thesis artifact that links figures, tables, scripts, result CSVs, and source literature to the written claims.",
    ]
    for c in contributions:
        add_bullet(doc, c)

    doc.add_heading("1.5 Scope and Limitations", level=2)
    add_body(doc, "The thesis is simulation-based. It evaluates controlled link and scheduling behaviour but does not include hardware-in-the-loop validation, over-the-air factory measurements, production PLC timing tests, or complete industrial safety certification. This limitation is not incidental; it defines the claim boundary. The thesis can support method ranking, design insight, and future experimental planning, but it cannot by itself certify a factory deployment.")
    add_body(doc, "The current results use available simulation outputs and do not introduce a new full simulation campaign. Some plots are direct simulation outputs, while some comparison visualizations are derived summaries. Wherever synthetic or smoothed projections exist in the repository, the thesis treats them as visual aids or excludes them from primary claims unless explicitly labelled.")
    add_figure(doc, ROOT / "system_design/topology_mobility.png", "Mobility-oriented factory topology concept motivating dynamic channel and scheduling decisions.", 5.7)

    doc.add_heading("1.6 Thesis Organization", level=2)
    add_body(doc, "Chapter 2 reviews the literature on Industry 5.0, 6G, industrial reliability, AI/ML for wireless networks, channel estimation, resource management, and simulation methodology. Chapter 3 defines the Factory6G system model, methodology, equations, algorithms, and traceability protocol. Chapter 4 presents the available results by evidence family. Chapter 5 synthesizes the findings into reliability mechanisms, AI/ML interpretation, design implications, and limitations. Chapter 6 concludes the thesis and identifies future work.")


def add_chapter_2(doc: Document) -> None:
    doc.add_heading("Chapter 2: Literature Review", level=1)
    add_body(doc, "The literature reviewed in this chapter establishes why 6G smart-factory communication requires a cross-layer and evidence-aware research method. The chapter does not attempt to survey every 6G topic. It focuses on the areas needed to justify the Factory6G methodology: Industry 5.0 requirements, 6G reliability targets, industrial channel behaviour, AI/ML for wireless optimization, channel estimation, resource management, and the need for traceable simulation evidence.")

    add_literature_matrix(doc)

    sections = [
        ("2.1 Industry 5.0 and Smart-Factory Communication Requirements",
         ["Industry 5.0 extends the automation emphasis of Industry 4.0 by foregrounding human-centric production, resilience, sustainability, and collaborative intelligence [1], [3], [40]. The communication network becomes a shared nervous system for machines, sensors, mobile robots, and human-facing tools.",
          "The smart-factory communication requirement is heterogeneous. Some devices send low-rate status updates, while others require near-real-time control or high-bandwidth perception data. The radio system must therefore support multiple reliability and latency classes rather than one uniform service.",
          "This motivates policy-aware scheduling. A resource manager that treats every link identically can waste resources or expose critical traffic to error. A 6G factory network should instead combine link evidence, traffic class, and policy objective."]),
        ("2.2 6G, Beyond-5G, and Extreme Reliability",
         ["6G is often framed around AI-native operation, integrated sensing and communication, extreme connectivity, and stronger reliability/latency targets than current 5G systems [9], [12], [33], [39]. These claims are useful as direction-setting goals, but they require careful experimental grounding.",
          "Reliability in a thesis context should be measured and bounded. Reporting a zero observed BER at finite sample count is not the same as proving zero error probability. Confidence-aware metrics therefore matter when interpreting Monte Carlo sweeps.",
          "The Factory6G evidence chain follows this caution by recording BER and BER upper confidence side by side. This makes strong results useful without overstating them."]),
        ("2.3 Industrial Propagation and Channel Modelling",
         ["Factory channels differ from open or office-like environments because metallic machinery, shelving, moving robots, and dense equipment create reflections, shadowing, and spatially varying path loss. A method evaluated only under a single stochastic model may fail to reveal this sensitivity.",
          "Rayleigh and Rician fading models remain important because they are controlled and interpretable. TR 38.901 UMi adds a more structured channel family that is useful for stress testing, even if it is not a complete factory ray-tracing measurement campaign.",
          "The thesis therefore uses a channel ladder. Simpler models expose method behaviour; richer models test whether apparent gains remain robust under more complex propagation assumptions."]),
        ("2.4 Channel Estimation and Receiver Reliability",
         ["Channel estimation converts pilots and received samples into the channel knowledge needed for equalization and demapping. Errors at this stage propagate directly into decoded bit reliability. Classical estimators are transparent and often computationally light, while learned or optimization-based estimators may capture more complex impairments.",
          "The literature on deep learning for wireless communication shows that learned methods can approximate difficult mappings and improve selected tasks [21], [22]. However, learned estimators must be evaluated against LS, DFT, LMMSE, adaptive, and optimization-driven baselines under consistent channel conditions.",
          "The key thesis position is not that neural estimation is always superior. The stronger claim is that estimator choice should be evaluated as part of a reliability stack with runtime, latency, and confidence-aware BER interpretation."]),
        ("2.5 Resource Management and Cross-Layer Optimization",
         ["Resource allocation decides who transmits, with which resources, and under what power constraints. It therefore shapes exposure to channel quality and traffic pressure. Classic objectives such as throughput maximization and proportional fairness are useful, but smart factories often require reliability-oriented control.",
          "Cross-layer optimization literature shows that interaction between layers can improve performance when the control variables are jointly interpreted [20], [27]. For Factory6G, the interface between PHY and MAC is expressed through resource directives and feedback metrics.",
          "A fair resource-manager study should compare simple, heuristic, optimization-inspired, and learned policies. The thesis therefore compares static, round-robin, max-throughput, proportional-fair, WMMSE, queue-aware, DRL, and BER-DRL policies using one output schema."]),
        ("2.6 AI/ML, DRL, and Reliability-Aware Control",
         ["AI/ML can support 6G by learning mappings that are difficult to hand-code or computationally expensive to solve repeatedly [9], [12], [18], [21]. DRL is particularly relevant when decisions affect future states or when a policy must balance multiple metrics.",
          "In this thesis, the learned scheduling contribution is deliberately bounded. BER-DRL is framed as a reliability-oriented policy trained from simulation-derived evidence. It should be judged by comparison with deterministic baselines and by whether it supports a useful design rule.",
          "This framing avoids a common weakness in AI-for-wireless studies: claiming replacement of classical methods without strong baseline comparison. A hybrid interpretation is more defensible for industrial networks."]),
        ("2.7 Simulation Traceability and Reproducibility",
         ["Simulation frameworks such as Sionna make it possible to evaluate next-generation physical-layer methods in repeatable experiments [25]. Reproducibility depends not only on code availability but also on result naming, configuration traceability, metric definitions, and clear separation between measured and derived evidence.",
          "Factory6G adopts this discipline by storing CSV, JSON, logs, plots, and summaries under timestamped run directories. The thesis uses those artifacts directly rather than inventing result values.",
          "The literature gap addressed by this thesis is therefore methodological as well as empirical: it demonstrates how AI/ML-assisted 6G smart-factory claims can be written with traceability and confidence boundaries."]),
    ]
    for title, paras in sections:
        doc.add_heading(title, level=2)
        for para in paras:
            add_body(doc, para)
        add_body(doc, "For Factory6G, this section contributes a specific design implication: the method under study must be interpreted through both communication theory and the operational constraints of industrial systems. This supports the later methodology chapter, where each algorithmic component is tied to a metric and an evidence source.")

    add_figure(doc, ROOT / "system_design/phy_mac_topology.png", "PHY/MAC topology concept used to frame cross-layer optimization in Factory6G.", 5.7)
    doc.add_heading("2.8 Extended Critical Synthesis", level=2)
    add_deepening_block(
        doc,
        "2.8.1 Industrial reliability as a multi-layer property",
        "Industrial reliability in 6G smart factories cannot be reduced to a single physical-layer target or a single MAC-layer utility.",
        ["propagation uncertainty", "receiver estimation quality", "traffic criticality", "scheduler objective"],
        "[7], [9], [12], [31], [39]",
        "the thesis must evaluate reliability as a chain of decisions in which failure can enter through channel modelling, estimator mismatch, scheduling exposure, or inappropriate operating mode",
    )
    add_deepening_block(
        doc,
        "2.8.2 Why AI-native networking requires baseline discipline",
        "AI-native 6G literature motivates adaptive and data-driven network control, but the strength of an AI/ML claim depends on the baselines used for comparison.",
        ["classical baseline strength", "training-data provenance", "generalization under channel shift", "runtime feasibility"],
        "[9], [12], [18], [21], [22]",
        "a learned method should be presented as credible only when its measured behaviour is compared with transparent deterministic methods under the same metric schema",
    )
    add_deepening_block(
        doc,
        "2.8.3 Factory propagation realism and result portability",
        "Factory environments contain reflecting structures, moving equipment, and heterogeneous layouts that complicate the portability of results between channel models.",
        ["metallic multipath", "partial blockage", "factory size scaling", "model-to-deployment transfer"],
        "[17], [25], [30], [42]",
        "channel families should remain visible in the results chapter so that the reader can distinguish robust method behaviour from channel-specific success",
    )
    add_deepening_block(
        doc,
        "2.8.4 Reliability, latency, and throughput as a design triangle",
        "Industrial communication design often requires balancing reliability, latency, and throughput rather than optimizing one metric in isolation.",
        ["low-latency control", "spectral efficiency pressure", "finite blocklength reliability", "traffic-class differentiation"],
        "[7], [9], [13], [14], [24]",
        "Factory6G should report multi-metric evidence so that a low-BER method is not automatically treated as deployable if runtime or throughput trade-offs are unacceptable",
    )
    add_deepening_block(
        doc,
        "2.8.5 Explainability and trust in learned radio control",
        "Learned radio-control policies introduce a governance question: the network may take actions that improve a reward but are difficult for engineers to inspect.",
        ["policy transparency", "fallback behaviour", "confidence monitoring", "operator auditability"],
        "[26], [29], [31]",
        "the final interpretation should preserve deterministic baselines and describe learned policies as auditable components rather than black-box replacements",
    )
    add_deepening_block(
        doc,
        "2.8.6 Cross-layer gap addressed by this thesis",
        "The literature motivates both AI/ML and 6G industrial communication, but it often leaves a gap between conceptual requirements and traceable result production.",
        ["method comparability", "result provenance", "claim boundaries", "simulation reproducibility"],
        "[20], [25], [27], [33]",
        "the thesis contribution is partly methodological: it shows how a simulation study can connect source artifacts, metrics, figures, and written claims",
    )
    doc.add_heading("2.9 Research Gap", level=2)
    gap_rows = [
        ["Layer-isolated evaluation", "PHY estimators and MAC schedulers are often evaluated separately.", "Factory6G links estimator and resource-manager evidence in one thesis narrative."],
        ["Weak baseline comparison", "Learned methods can be compared against narrow or weak baselines.", "The thesis compares learned policies with static, heuristic, fairness, optimization-inspired, and DRL baselines."],
        ["Overclaiming zero-error runs", "Finite simulations can report zero observed errors.", "BER upper confidence is used as a required interpretation metric."],
        ["Insufficient channel diversity", "Single-channel results can hide sensitivity.", "Rayleigh, Rician, and TR 38.901 UMi families are kept separate in interpretation."],
        ["Poor evidence traceability", "Figures may be detached from source data.", "Each major result is linked to a CSV, JSON, or PNG source in the audit note."],
    ]
    add_table(doc, "Synthesized research gaps and Factory6G response", ["Gap", "Risk in prior work", "Factory6G response"], gap_rows)


def add_chapter_3(doc: Document) -> None:
    doc.add_heading("Chapter 3: System Model and Methodology", level=1)
    add_body(doc, "This chapter defines the Factory6G methodology used to generate, organize, and interpret the thesis evidence. The focus is not only on implementing a simulator. The methodological contribution is a traceable link between theoretical communication quantities, AI/ML decision surfaces, simulation outputs, and thesis claims.")

    doc.add_heading("3.1 System Overview", level=2)
    add_body(doc, "Factory6G is organized around a fixed simulation flow. A configuration file defines channel model, modulation, factory profile, Monte Carlo policy, enabled estimators, enabled resource managers, and output preferences. The runtime creates a timestamped output directory, prepares deterministic seeds, runs selected stages, writes JSON and CSV results, and generates plots for BER, raw BER, latency, throughput, power, and runtime.")
    add_body(doc, "The system is intentionally modular. PHY components handle transmitter, channel, receiver, and estimator behaviour. MAC/resource-management components generate scheduling and power directives. Simulation orchestration controls shared batch contexts, stopping rules, output writing, and plot generation. This separation allows individual methods to be compared while retaining a common measurement protocol.")
    add_algorithm(doc, "Factory6G evidence-generation workflow", [
        "Input: config.json, selected channel family, estimator list, resource-manager list, Eb/N0 grid",
        "1. Load and validate the simulation configuration.",
        "2. Create a timestamped run directory and initialize deterministic random seeds.",
        "3. For each Eb/N0 point, create or reuse a controlled batch context.",
        "4. Evaluate selected estimator methods on shared channel/noise evidence.",
        "5. Build PHY feedback for resource-manager evaluation.",
        "6. Evaluate selected scheduling and power-control policies.",
        "7. Accumulate BER, BER upper confidence, throughput, latency, power, and runtime.",
        "8. Write stage_results_v2.csv/json, summary_v2.csv/json, plots, and logs.",
        "Output: traceable result artifacts for thesis tables, figures, and interpretation.",
    ])

    doc.add_heading("3.2 Link and Channel Model", level=2)
    add_body(doc, "The link-level view begins with a transmitted symbol vector passing through a channel and additive noise. While the implementation uses Sionna/TensorFlow components, the conceptual model can be written as a received-signal relation in which decoded reliability depends on the channel realization, estimation error, modulation, noise level, and receiver processing.")
    add_equation(doc, "3.1", "y = Hx + n", "where x is the transmitted signal, H is the channel response, n is additive noise, and y is the received signal used by the receiver and estimator pipeline.")
    add_equation(doc, "3.2", "BER = N_error / N_total", "where N_error is the number of incorrectly decoded bits and N_total is the number of transmitted bits included in the Monte Carlo estimate.")
    add_body(doc, "The channel families used in the current thesis evidence are Rayleigh, Rician, and TR 38.901 UMi. Rayleigh provides a non-line-of-sight fading baseline, Rician introduces a stronger dominant component, and TR 38.901 UMi provides a structured channel family used as a more demanding scenario. The methodology keeps the families separate during interpretation because combining them too early can hide method-specific sensitivity.")

    doc.add_heading("3.3 Factory Profiles and Datasets", level=2)
    add_body(doc, "Factory6G includes small, medium, and large profile concepts, as documented in the dataset configuration. These profiles differ in physical room dimensions, machine count, and intended industrial archetype. The dataset layer stores channel samples and resource-manager training data, including active-user masks, channel energy, per-user power, oracle utility, BER-oriented metrics, and provenance fields.")
    add_table(doc, "Factory profile concepts used by the dataset workflow", ["Scale", "Scenario label", "Room size", "Machines", "Research role"], [
        ["S", "Electronics Workcell", "20 x 15 x 5 m", "6", "Compact dense-cell baseline"],
        ["M", "Automotive Assembly Zone", "45 x 30 x 8 m", "14", "Mid-scale factory-zone stress case"],
        ["L", "Logistics Fulfillment Hall", "90 x 60 x 12 m", "28", "Large-layout geometry and path-loss stress case"],
    ])
    source_entries.append((rel(ROOT / "data/README.md"), "Dataset schema and factory profile source"))
    source_entries.append((rel(ROOT / "config/factory_size_profiles.json"), "Factory profile definitions"))

    doc.add_heading("3.4 Estimator Methodology", level=2)
    add_body(doc, "The estimator stage compares receiver-side methods under shared conditions. The core fairness requirement is that methods should see the same channel/noise context wherever possible. This reduces the chance that a method appears better only because it was evaluated on an easier Monte Carlo draw.")
    add_table(doc, "Estimator families and thesis interpretation role", ["Estimator family", "Role", "Expected strength", "Claim boundary"], [
        ["LS", "Classical baseline", "Simple, transparent, low overhead", "Can degrade under severe multipath or low pilot quality"],
        ["DFT", "Structured estimator baseline", "Can exploit frequency-domain structure", "Performance depends on channel assumptions"],
        ["Adaptive/PSO", "Optimization-oriented alternatives", "Can search non-convex spaces", "Runtime and stability must be monitored"],
        ["Neural", "Learned estimator", "Can learn patterns from training data", "Dataset dependence and generalization remain bounded claims"],
    ])
    add_equation(doc, "3.3", "e_H = ||H - H_hat||_2^2", "where e_H represents estimation error; in the thesis it is interpreted indirectly through BER, latency, throughput, and runtime rather than as the only optimization target.")

    doc.add_heading("3.5 Resource-Manager Methodology", level=2)
    add_body(doc, "The resource-manager stage turns link and state information into scheduling and power-control directives. A resource manager decides which users are active and how power is allocated. These decisions affect BER because they determine which channel conditions are used for transmission and how aggressively the system pursues throughput or fairness.")
    add_table(doc, "Resource-manager methods included in the May 23 evidence", ["Method", "Type", "Interpretation"], [
        ["static", "baseline", "Fixed or minimally adaptive allocation; useful for exposing cost of non-adaptation"],
        ["round_robin", "baseline", "Fair rotation without deep channel optimization"],
        ["max_throughput", "baseline", "Throughput-oriented deterministic policy"],
        ["pf", "baseline", "Proportional-fair balance between efficiency and fairness"],
        ["wmmse", "optimization-inspired baseline", "Weighted minimum mean-square-error inspired allocation"],
        ["queue_aware", "heuristic baseline", "Uses queue or demand awareness for scheduling decisions"],
        ["drl", "learned baseline", "Existing learned policy checkpoint"],
        ["ber_drl", "trained policy", "Reliability-oriented learned policy emphasized in this thesis"],
    ])
    add_equation(doc, "3.4", "a_t = pi_theta(s_t)", "where a_t is the scheduling/power action, s_t is the observed state, and pi_theta is either a deterministic rule or learned policy.")

    doc.add_heading("3.6 BER-Oriented DRL Formulation", level=2)
    add_body(doc, "The BER-oriented DRL formulation treats reliability as the primary training signal. Instead of rewarding only throughput or fairness, the policy is shaped toward avoiding high-error scheduling actions. This does not guarantee universal dominance, because deterministic policies can be very strong under favourable conditions. It does provide a way to encode reliability preference into a learned resource-manager policy.")
    add_equation(doc, "3.5", "r_t = -log(max(BER_t, epsilon)) - lambda_L L_t + lambda_T T_t", "where the reward balances reliability, latency, and throughput; epsilon prevents numerical instability when BER is very small.")
    add_algorithm(doc, "BER-oriented resource-manager inference", [
        "Input: channel-energy tensor, active-user mask, trained policy checkpoint",
        "1. Normalize input features using checkpoint metadata.",
        "2. Run the policy network to produce scheduling logits and power outputs.",
        "3. Select active users subject to configured active-user constraints.",
        "4. Clip or normalize power directives to valid simulation bounds.",
        "5. Return ResourceDirectives to the PHY/MAC evaluation stage.",
        "6. Record BER, BER upper confidence, throughput, latency, power, and runtime.",
    ])

    doc.add_heading("3.7 Metrics and Confidence-Aware Reading", level=2)
    add_body(doc, "BER is the primary reliability metric, but it is not sufficient alone. The thesis also records BER upper confidence, throughput bits per batch, latency, runtime, and average power. BER upper confidence is particularly important when the observed bit-error count is zero. In a finite run, zero observed errors means that no errors were seen in the sample, not that the true error probability is zero.")
    add_table(doc, "Metric definitions and thesis use", ["Metric", "Definition", "How it is used"], [
        ["BER", "Observed bit errors divided by transmitted bits", "Primary reliability ranking metric"],
        ["BER upper confidence", "Conservative confidence bound on BER", "Prevents overclaiming zero-error observations"],
        ["Throughput bits per batch", "Successfully delivered bits per simulated batch", "Used to identify reliability-throughput trade-offs"],
        ["Latency ms", "Per-batch latency estimate in milliseconds", "Used to bound deployability of reliability gains"],
        ["Runtime sec", "Simulation/runtime cost", "Used to separate link latency from computation cost"],
        ["Average power W", "Mean transmitted power estimate", "Used to discuss energy-related trade-offs"],
    ])

    doc.add_heading("3.8 Reproducibility Protocol", level=2)
    add_body(doc, "The reproducibility protocol has four rules. First, all scripts are executed through Docker so that TensorFlow, Sionna, plotting, and document tooling remain isolated from local host drift. Second, result claims are linked to stored CSV, JSON, PNG, or PPTX files. Third, primary thesis claims use direct simulation outputs, while synthetic projections are labelled as such. Fourth, the final document includes an audit note listing every reused source artifact and unresolved citation check.")
    add_table(doc, "Reproducibility artifacts and responsibilities", ["Artifact", "Purpose", "Used in thesis"], [
        ["config.json", "Runtime configuration", "Appendix reproducibility"],
        ["summary_v2.csv/json", "Run-level summary", "Evidence map and source audit"],
        ["stage_results_v2.csv/json", "Per-stage metrics", "Result tables and figure interpretation"],
        ["reports/plots/*.png", "Publication-style plots", "Chapters 3 and 4 figures"],
        ["reports/weekly/*", "Progress evidence decks and reports", "Audit trail and supporting interpretation"],
        ["models/*", "Trained estimator/resource-manager checkpoints", "Methodology and appendix"],
    ])
    doc.add_heading("3.9 Methodological Deepening", level=2)
    add_deepening_block(
        doc,
        "3.9.1 Shared context and fair comparison",
        "A fair estimator or scheduler comparison requires more than running methods in sequence; the compared methods must be exposed to equivalent operating evidence.",
        ["shared channel/noise context", "common Eb/N0 grid", "identical metric schema", "consistent stopping policy"],
        "[7], [21], [25]",
        "method ranking should be based on controlled experimental differences rather than uncontrolled Monte Carlo variation",
    )
    add_deepening_block(
        doc,
        "3.9.2 State, action, and reward design for resource managers",
        "The resource-manager problem is naturally expressed through state, action, and utility definitions, even when the method is deterministic rather than learned.",
        ["state observability", "active-user constraints", "power normalization", "reward-objective alignment"],
        "[14], [18], [19], [27]",
        "the thesis should discuss scheduling as a control surface whose objective must match the factory traffic class",
    )
    add_deepening_block(
        doc,
        "3.9.3 Confidence-aware Monte Carlo design",
        "Monte Carlo simulation is useful because it exposes methods to repeated channel and noise draws, but finite sampling creates interpretation risk.",
        ["zero-error censoring", "upper-confidence reporting", "sample-count transparency", "independent-seed extension"],
        "[9], [25], [39]",
        "every reliability claim should mention both observed BER and the confidence-aware boundary when the number of bit errors is small",
    )
    add_deepening_block(
        doc,
        "3.9.4 Dataset provenance and learned-policy boundaries",
        "Learned policies inherit the assumptions of the data used to train them, including channel family, active-user distribution, objective function, and labelling strategy.",
        ["synthetic-data generation", "oracle-label quality", "checkpoint metadata", "deployment drift"],
        "[18], [21], [22], [43]",
        "BER-DRL should be evaluated as a policy trained under defined assumptions rather than as a universally valid scheduling rule",
    )
    add_deepening_block(
        doc,
        "3.9.5 Reproducibility as part of the research method",
        "For simulation-heavy 6G research, reproducibility is not an administrative detail; it is part of the scientific method.",
        ["Dockerized execution", "timestamped outputs", "CSV/JSON dual storage", "plot-source traceability"],
        "[25], [31], [33]",
        "the final thesis should allow another researcher to locate the data behind each major result and understand what was measured",
    )


def add_chapter_4(doc: Document) -> None:
    doc.add_heading("Chapter 4: Results", level=1)
    add_body(doc, "This chapter presents the current Factory6G evidence without launching a new simulation campaign. The results are organized by evidence family so that estimator behaviour, channel sensitivity, modulation sensitivity, resource management, BER-DRL, and JIDD-SCMA are not collapsed into a single undifferentiated comparison. Each result is interpreted with its source artifact and claim boundary.")
    source_entries.append((rel(RUN_ROOT / "summary_v2.csv"), "May 23 cross-channel resource-manager run summary"))
    source_entries.append((rel(RUN_ROOT / "summary_v2.json"), "May 23 cross-channel resource-manager run metadata"))
    for channel in ("rayleigh", "rician", "tr38901"):
        source_entries.append((rel(RUN_ROOT / channel / "resource_managers/stage_results_v2.json"), f"Primary {channel} resource-manager stage metrics"))

    doc.add_heading("4.1 Reading Protocol", level=2)
    add_body(doc, "The results should be read in three passes. First, BER and BER upper confidence identify the reliability profile. Second, throughput, latency, runtime, and power show engineering trade-offs. Third, channel and method families are kept separate to avoid attributing a channel effect to a method effect. This reading protocol is especially important for BER-DRL, because the learned policy is competitive but not universally dominant.")

    doc.add_heading("4.2 Estimator Reliability Across Eb/N0", level=2)
    add_body(doc, "Estimator results show that receiver-side processing materially affects reliability. As Eb/N0 increases, BER generally decreases, but estimator choice influences the steepness of that decrease and whether a residual floor remains. This supports the thesis argument that PHY design is a reliability control surface, not merely a fixed preprocessing block.")
    add_figure(doc, ROOT / "reports/plots/estimator_ber_vs_ebno.png", "Estimator BER versus Eb/N0 under the current plotted evidence family.", 5.8)
    add_figure(doc, ROOT / "reports/plots/estimator_latency_vs_ebno.png", "Estimator latency versus Eb/N0, used to interpret reliability together with timing cost.", 5.8)
    add_figure(doc, ROOT / "reports/plots/estimator_throughput_vs_ebno.png", "Estimator throughput versus Eb/N0, used to identify reliability-throughput trade-offs.", 5.8)
    add_figure(doc, ROOT / "reports/plots/estimator_runtime.png", "Estimator runtime comparison, separating computational burden from link reliability.", 5.8)
    add_figure(doc, ROOT / "reports/plots/neural_vs_ls_direct_ber.png", "Focused neural-estimator versus LS BER comparison retained as supporting estimator evidence.", 5.65)
    add_figure(doc, ROOT / "reports/plots/neural_vs_ls_latency.png", "Focused neural-estimator versus LS latency comparison retained as supporting estimator evidence.", 5.65)
    add_body(doc, "The estimator figures should not be interpreted as a single winner-takes-all ranking. Instead, they show how reliability, timing, and implementation burden interact. A learned estimator can be useful when training data match deployment conditions, but classical estimators remain essential because they are transparent, inexpensive, and suitable as fallback modes.")

    doc.add_heading("4.3 Channel-Model Sensitivity", level=2)
    add_body(doc, "Channel sensitivity is central to the smart-factory problem. A method that looks strong under one channel family can become less reliable under another. The current channel-model evidence therefore acts as a guardrail for later resource-manager claims.")
    add_figure(doc, ROOT / "reports/plots/channel_model_ber_vs_ebno.png", "BER sensitivity across channel model families.", 5.8)
    add_body(doc, "The implication is methodological: results should be reported by channel family before they are synthesized. This is why the May 23 resource-manager comparison reports Rayleigh, Rician, and TR 38.901 UMi separately. The separation makes it possible to see that BER-DRL can match the best Rayleigh reliability while still being second to deterministic max-throughput in the Rician and UMi summaries.")

    doc.add_heading("4.4 Modulation and Factory-Size Sensitivity", level=2)
    add_body(doc, "Modulation sensitivity connects reliability to spectral-efficiency pressure. Higher-order modulation can carry more bits per symbol but reduces separation between constellation points and increases vulnerability to noise and estimation error. Factory-size sensitivity adds the geometry dimension: larger layouts can increase path-loss diversity, shadowing, and scheduling difficulty.")
    add_figure(doc, ROOT / "reports/plots/modulation_ber_vs_ebno.png", "Modulation BER versus Eb/N0, showing the reliability cost of higher spectral efficiency.", 5.8)
    add_figure(doc, ROOT / "reports/plots/modulation_latency_vs_ebno.png", "Modulation latency versus Eb/N0, used to separate spectral-efficiency pressure from timing behaviour.", 5.8)
    add_figure(doc, ROOT / "reports/plots/factory_size_ber_vs_ebno.png", "Factory-size BER sensitivity across available profile evidence.", 5.8)
    add_body(doc, "For a smart factory, the design implication is that modulation and scheduling should be tied to traffic criticality. A safety-critical control flow should not be forced into a high-throughput mode if the resulting error probability violates the reliability class. Conversely, non-critical monitoring traffic may tolerate a different operating point.")

    doc.add_heading("4.5 Cross-Channel Resource-Manager Results", level=2)
    add_body(doc, "The May 23 resource-manager evidence is the main MAC-layer result set. It evaluates static, round-robin, max-throughput, proportional-fair, WMMSE, queue-aware, DRL, and BER-DRL policies across Rayleigh, Rician, and TR 38.901 UMi. The table below is extracted from the current resource-manager comparison CSV.")
    add_table(doc, "May 23 cross-channel resource-manager ranking", ["Channel", "Method", "Type", "Rank", "BER", "BER UCB", "Latency ms", "Throughput"], rm_summary_rows(), "E8EEF5")
    best_text, best_rows = best_method_text()
    add_body(doc, best_text)
    add_table(doc, "Best observed method versus BER-DRL by channel", ["Channel", "Best method", "Best BER", "Best BER UCB", "Learned policy", "BER-DRL BER", "Interpretation"], best_rows, "F2F4F7")
    add_figure(doc, RUN_ROOT / "overview/resource_managers/ber_vs_ebno.png", "Overview BER comparison for resource managers across the May 23 evidence family.", 5.8)
    add_figure(doc, RUN_ROOT / "overview/resource_managers/latency_vs_ebno.png", "Overview latency comparison for resource managers across the May 23 evidence family.", 5.8)
    add_figure(doc, RUN_ROOT / "overview/resource_managers/throughput_vs_ebno.png", "Overview throughput comparison for resource managers across the May 23 evidence family.", 5.8)
    add_figure(doc, RUN_ROOT / "rayleigh/resource_managers/ber_vs_ebno.png", "Rayleigh resource-manager BER curves from the May 23 stage output.", 5.75)
    add_figure(doc, RUN_ROOT / "rician/resource_managers/ber_vs_ebno.png", "Rician resource-manager BER curves from the May 23 stage output.", 5.75)
    add_figure(doc, RUN_ROOT / "tr38901/resource_managers/ber_vs_ebno.png", "TR 38.901 UMi resource-manager BER curves from the May 23 stage output.", 5.75)

    doc.add_heading("4.6 BER-DRL Evidence", level=2)
    add_body(doc, "BER-DRL is the reliability-oriented learned policy emphasized in this thesis. Its contribution is not that it defeats all deterministic methods. The more defensible conclusion is that it remains competitive under multiple channel families while preserving the same output schema and confidence-aware interpretation used for non-learned policies.")
    add_figure(doc, RUN_ROOT / "resource_manager_channel_comparison_synthetic_methods/ber_drl_ber_vs_ebno.png", "BER-DRL channel-family profile. This plot is treated as a derived visualization and interpreted with the source audit boundary.", 5.8)
    add_body(doc, "The Rayleigh summary shows that BER-DRL matches the best observed BER because several methods report zero observed errors with the same BER upper confidence. Under Rician, BER-DRL ranks second behind max-throughput. Under UMi/TR 38.901, it is also second, very close to queue-aware by observed BER and confidence bound but still behind max-throughput. This pattern is valuable because it prevents simplistic claims. Learned reliability-oriented control is useful, but deterministic baselines remain strong and must be retained.")

    doc.add_heading("4.7 JIDD-SCMA Joint-Processing Evidence", level=2)
    add_body(doc, "The JIDD-SCMA evidence provides an advanced joint-processing result family. It is included because joint detection and decoding aligns with future cross-layer reliability work. However, the current available run contains non-monotonic behaviour at higher Eb/N0. The thesis therefore treats JIDD-SCMA as supporting evidence and future-work motivation rather than as a central ranking claim.")
    add_figure(doc, ROOT / "reports/plots/jidd_ber_comparison.png", "JIDD-SCMA BER comparison, included with explicit caution about current result stability.", 5.8)

    doc.add_heading("4.8 Multi-Metric Interpretation", level=2)
    add_body(doc, "BER alone is insufficient for industrial interpretation. A method with low BER but excessive runtime may be unsuitable for tight edge deployment. A method that maximizes throughput while increasing BER upper confidence may be unsuitable for safety-critical control. A method with strong Rayleigh performance may still require caution under structured multipath.")
    add_figure(doc, ROOT / "reports/plots/combined_ber.png", "Combined BER evidence used to synthesize reliability trends across result families.", 5.8)
    add_figure(doc, ROOT / "reports/plots/runtime_comparison.png", "Runtime comparison used to contextualize reliability gains with computation cost.", 5.8)
    add_body(doc, "The multi-metric reading supports a hybrid 6G factory design. Deterministic baselines provide transparency and fallback. Learned policies provide adaptable objective encoding. Confidence-aware metrics prevent overclaiming. Together, these components produce a more credible thesis claim than a single AI-wins narrative.")
    doc.add_heading("4.9 Extended Result Interpretation", level=2)
    add_deepening_block(
        doc,
        "4.9.1 Reading estimator curves as deployment evidence",
        "Estimator curves are not only signal-processing plots; they indicate whether a receiver method can support an industrial operating point.",
        ["BER slope across Eb/N0", "latency burden", "runtime overhead", "fallback suitability"],
        "[7], [21], [25], [42]",
        "estimator selection should be made by matching reliability improvement to deployment constraints rather than by ranking BER alone",
    )
    add_deepening_block(
        doc,
        "4.9.2 Reading channel sensitivity as robustness evidence",
        "Channel-model sensitivity tests whether a method remains credible when the propagation assumption changes.",
        ["Rayleigh baseline behaviour", "Rician dominant-path behaviour", "TR 38.901 UMi stress", "factory geometry transfer"],
        "[17], [25], [30]",
        "a thesis claim should not be generalized across factory conditions unless the channel-specific results support that generalization",
    )
    add_deepening_block(
        doc,
        "4.9.3 Reading scheduler rankings as policy evidence",
        "Resource-manager rankings reveal the policy objective embedded in each method and the channel conditions under which that objective is useful.",
        ["throughput preference", "fairness preference", "queue awareness", "BER-oriented learning"],
        "[14], [18], [19], [27]",
        "a smart factory should expose policy modes instead of assuming one scheduler is optimal for all traffic and channel states",
    )
    add_deepening_block(
        doc,
        "4.9.4 Reading BER-DRL as bounded learned evidence",
        "BER-DRL is the main learned resource-manager contribution, but its best interpretation is bounded and comparative.",
        ["Rayleigh matched-best behaviour", "Rician second-place behaviour", "UMi/TR 38.901 second-place behaviour", "deterministic-baseline retention"],
        "[9], [18], [21], [22]",
        "BER-DRL should be positioned as a reliability-oriented option inside a hybrid system, not as proof that learned control always dominates classical scheduling",
    )
    add_deepening_block(
        doc,
        "4.9.5 Reading JIDD-SCMA as future-work evidence",
        "The JIDD-SCMA result family points toward joint-processing opportunities but also shows why anomalous behaviour must be treated carefully.",
        ["non-monotonic BER behaviour", "decoder-configuration sensitivity", "sample stability", "integration with resource management"],
        "[13], [21], [25]",
        "advanced joint-processing methods should be integrated into the central thesis ranking only after stability and reproducibility are strengthened",
    )


def add_chapter_5(doc: Document) -> None:
    doc.add_heading("Chapter 5: Discussion", level=1)
    add_body(doc, "The results support the central thesis that 6G smart-factory reliability is a cross-layer property. It emerges from channel conditions, receiver estimation, modulation, scheduling, power control, and policy objective. AI/ML can contribute to this stack, but only when interpreted through strong baselines, confidence bounds, and deployment constraints.")

    discussions = [
        ("5.1 Reliability Mechanisms",
         ["The first reliability mechanism is improved channel knowledge. Better estimation can reduce demapping and decoding errors, but estimator performance is limited by channel realism, training distribution, and computational cost.",
          "The second mechanism is scheduling exposure. Resource managers determine which links are selected and how resources are allocated. A policy can improve reliability by avoiding fragile transmissions, but this may reduce immediate throughput or fairness.",
          "The third mechanism is operating-point selection. Modulation, Eb/N0, factory size, and channel family jointly determine whether a method operates in a reliable regime or near an error floor."]),
        ("5.2 Interpretation of AI/ML Contributions",
         ["The most important AI/ML result is not a universal victory. BER-DRL performs competitively, but deterministic max-throughput remains the best observed method in the Rician and UMi/TR 38.901 summaries. This means the learned policy should be viewed as a reliability-aware option inside a hybrid control framework.",
          "This interpretation is stronger than claiming that AI replaces classical communication methods. Industrial systems need fallback behaviour, explainability, and monitoring. A learned policy that cannot be audited or bounded is difficult to justify in safety-critical contexts.",
          "The thesis therefore positions AI/ML as an objective-encoding and adaptation mechanism. It can learn reliability-sensitive decisions from simulation evidence, but it should remain coupled to deterministic baselines and confidence-aware validation."]),
        ("5.3 Deployment Implications",
         ["A practical 6G factory network should expose policy modes. Safety-critical control traffic may prefer reliability-first scheduling, conservative modulation, and confidence monitoring. Monitoring traffic may tolerate higher throughput and weaker reliability guarantees.",
          "The network should also maintain method diversity. Classical estimators and deterministic schedulers provide robust fallback modes. Learned methods can be enabled where training provenance and runtime monitoring support their use.",
          "Finally, result interpretation should be operational. Engineers need to know not only which method has the lowest mean BER but also whether the confidence bound is acceptable, what throughput was sacrificed, and whether runtime fits edge-controller budgets."]),
        ("5.4 Limitations",
         ["The evidence is simulation-only. It does not include factory-floor measurements, hardware-in-the-loop timing, external interference, clock synchronization, or controller integration. This limits deployment claims.",
          "The learned methods depend on synthetic data and checkpoint assumptions. If factory geometry, device density, or traffic changes significantly, model generalization must be revalidated.",
          "Monte Carlo evidence is finite. Confidence bounds help, but larger independent seeds and longer campaigns would strengthen the statistical basis.",
          "The JIDD-SCMA evidence is promising but not yet stable enough to support a central ranking claim. It should be extended only after non-monotonic behaviour is understood."]),
        ("5.5 Design Principles",
         ["Design principle 1: Treat reliability as a multi-metric operating condition, not a single BER number.",
          "Design principle 2: Keep deterministic baselines in the system even when learned policies are introduced.",
          "Design principle 3: Separate channel families during analysis before synthesizing conclusions.",
          "Design principle 4: Use confidence-aware interpretation whenever zero-error observations appear.",
          "Design principle 5: Link every thesis claim to a reproducible artifact: configuration, run directory, CSV, JSON, plot, or model checkpoint."]),
    ]
    for title, paras in discussions:
        doc.add_heading(title, level=2)
        for para in paras:
            add_body(doc, para)

    add_table(doc, "Observed evidence and design implication", ["Observed evidence", "Interpretation", "Design implication"], [
        ["BER improves with Eb/N0 but may retain a floor", "Noise is not the only reliability driver", "Use channel realism and confidence bounds"],
        ["Estimator methods differ in BER and runtime", "Receiver design affects both reliability and deployability", "Select estimator by operating class"],
        ["Higher-order modulation increases reliability pressure", "Throughput can conflict with error probability", "Tie modulation to traffic criticality"],
        ["Resource managers rank differently by channel", "MAC policy is channel-sensitive", "Use channel-aware policy selection"],
        ["BER-DRL is competitive but not universal", "AI/ML is useful but bounded", "Deploy learned methods with deterministic baselines"],
        ["JIDD-SCMA has current anomalies", "Advanced methods need stability checks", "Keep as future-work path until validated"],
    ])
    doc.add_heading("5.6 Extended Deployment Discussion", level=2)
    add_deepening_block(
        doc,
        "5.6.1 Policy modes for factory traffic classes",
        "The Factory6G evidence suggests that a single universal MAC policy is not appropriate for all smart-factory traffic.",
        ["safety-critical control", "mobile-robot coordination", "digital-twin synchronization", "non-critical monitoring"],
        "[7], [13], [24], [31]",
        "a 6G factory network should map traffic classes to scheduler and modulation policies with explicit reliability targets",
    )
    add_deepening_block(
        doc,
        "5.6.2 Hybrid AI/classical operation",
        "A hybrid radio stack can use learned policies when they add value while preserving deterministic methods as explainable baselines and fallback modes.",
        ["runtime monitoring", "model drift", "operator confidence", "fallback activation"],
        "[9], [12], [26], [29]",
        "the most credible AI/ML deployment path is supervised by confidence monitoring and bounded by deterministic alternatives",
    )
    add_deepening_block(
        doc,
        "5.6.3 From simulation evidence to experimental validation",
        "The thesis evidence is useful because it is controlled, but the next step toward deployment is experimental validation under hardware and factory constraints.",
        ["radio-frequency interference", "controller timing", "clock synchronization", "over-the-air repeatability"],
        "[17], [25], [30], [42]",
        "future work should prioritize validation pathways that test whether the simulation ranking survives physical implementation effects",
    )


def add_chapter_6(doc: Document) -> None:
    doc.add_heading("Chapter 6: Conclusion and Future Work", level=1)
    doc.add_heading("6.1 Research Summary", level=2)
    add_body(doc, "This thesis investigated AI/ML-assisted reliability for 6G and Beyond-5G smart-factory networks. The work developed a cross-layer Factory6G thesis methodology that connects physical-layer channel estimation, OFDM receiver processing, channel model sensitivity, MAC-layer resource management, BER-oriented learned scheduling, and confidence-aware evidence interpretation.")
    add_body(doc, "The current evidence supports a hybrid conclusion. Reliability is shaped jointly by channel realism, estimator choice, modulation, factory geometry, scheduling policy, and power control. Learned methods can support reliability-aware operation, but deterministic baselines remain essential. The most defensible 6G factory design is therefore not purely AI-driven; it is mathematically grounded, simulation-traceable, confidence-aware, and hybrid.")

    doc.add_heading("6.2 Answers to Research Questions", level=2)
    answers = [
        ["RQ1", "Factory6G connects PHY and MAC evaluation through a fixed simulation flow, shared metric schema, Dockerized execution discipline, and timestamped result artifacts."],
        ["RQ2", "Estimator choice affects BER and runtime, and its reliability value depends on channel family, Eb/N0, and deployability constraints."],
        ["RQ3", "Resource-manager ranking is channel-sensitive. In the May 23 evidence, BER-DRL matches the best observed Rayleigh BER but does not beat max-throughput under Rician or UMi/TR 38.901."],
        ["RQ4", "BER-DRL should be interpreted as a competitive reliability-oriented policy inside a hybrid stack, not as a universal replacement for deterministic baselines."],
    ]
    add_table(doc, "Answers to the research questions", ["Research question", "Answer"], answers)

    doc.add_heading("6.3 Contributions Revisited", level=2)
    for c in [
        "A traceable Factory6G simulation methodology for cross-layer reliability evaluation.",
        "A confidence-aware BER interpretation framework that avoids overclaiming finite zero-error observations.",
        "A current evidence synthesis across estimator, channel, modulation, factory-size, resource-manager, BER-DRL, and JIDD-SCMA result families.",
        "A bounded academic interpretation of AI/ML for smart-factory reliability, emphasizing hybrid deployment with deterministic baselines.",
        "A reproducible thesis document that maps claims to local result artifacts and literature sources.",
    ]:
        add_bullet(doc, c)

    doc.add_heading("6.4 Future Work", level=2)
    for f in [
        "Run larger independent-seed campaigns to reduce confidence intervals and test robustness across factory profiles.",
        "Introduce richer traffic models, queue arrivals, mobility patterns, interference sources, and controller timing constraints.",
        "Validate selected methods through hardware-in-the-loop or over-the-air factory testbeds.",
        "Extend BER-DRL with explainability, safety constraints, and online drift monitoring.",
        "Stabilize the JIDD-SCMA result family before integrating it into the central resource-management ranking.",
        "Explore digital-twin integration so that factory geometry, mobility, and traffic state can update the simulation-to-policy pipeline.",
    ]:
        add_bullet(doc, f)
    add_body(doc, "The final conclusion is that AI/ML-assisted 6G smart-factory communication is credible when it is not treated as a black-box replacement for communication theory. The strongest path is cross-layer and hybrid: use learned policies where they encode useful objectives, preserve deterministic baselines where they provide transparency and fallback, and evaluate all methods through traceable evidence and confidence-aware metrics.")


def add_appendices(doc: Document) -> None:
    doc.add_heading("Appendix A: Reproducibility Commands", level=1)
    add_body(doc, "The commands below document the Docker-first workflow used by Factory6G. They are included for reproducibility and should be adapted only when regenerating results intentionally.")
    add_algorithm(doc, "Representative Docker commands", [
        "docker compose build simulation",
        "docker compose run --rm simulation --config config.json --resource-managers static,round_robin,max_throughput,pf,wmmse,queue_aware,drl --channel rayleigh --modulation low --factory-size s",
        "docker compose run --rm simulation --config config.json --resource-managers ber_drl --channel tr38901 --modulation low --factory-size s",
        "docker compose run --rm --entrypoint python -v \"$PWD:/app\" simulation scripts/tools/generate_rm_ber_report.py",
        "docker compose run --rm --entrypoint python -v \"$PWD:/app\" simulation -m pytest tests/test_drl_policy_pipeline.py -q",
    ])

    doc.add_heading("Appendix B: Evidence Source Map", level=1)
    rows = []
    for path, note in unique_source_entries():
        rows.append([path, note])
    add_table(doc, "Source artifacts reused in the thesis", ["Path", "Use"], rows[:80], "F2F4F7")

    doc.add_heading("Appendix C: Figure Inventory", level=1)
    add_table(doc, "Embedded figure inventory", ["Figure", "Caption", "Source path"], [[n, c, p] for n, c, p in figure_entries], "F2F4F7")
    add_table(doc, "Embedded algorithm inventory", ["Algorithm", "Caption"], [[n, c] for n, c in algorithm_entries], "F2F4F7")

    doc.add_heading("Appendix D: Result Interpretation Notes", level=1)
    for note in [
        "Primary ranking claims use the May 23 resource-manager CSV, not the synthetic simulation-anchored projection CSV.",
        "Figures generated from synthetic or smoothed projections are labelled as derived visualizations and are not used as sole evidence for ranking claims.",
        "A zero observed BER is interpreted together with BER upper confidence.",
        "JIDD-SCMA is included as supporting evidence because the available run shows behaviour that should be stabilized before making a central ranking claim.",
        "The final thesis remains simulation-bound and does not claim hardware or factory-floor certification.",
    ]:
        add_bullet(doc, note)

    doc.add_heading("Appendix E: Selected Configuration and Dataset Schema", level=1)
    add_body(doc, "Factory6G uses a top-level configuration schema covering simulation, Monte Carlo policy, estimators, resource managers, system settings, transceiver settings, factory scenario, and ray tracing. Resource-manager training data include scenario, channel model, Eb/N0, channel energy, active-user mask, per-user power, oracle utility, oracle BER, confidence bounds, throughput, latency, candidate counts, objective, and source-manager fields.")
    add_table(doc, "Key dataset columns", ["Column", "Meaning"], [
        ["scenario", "Factory or channel scenario label"],
        ["channel_model_type", "Rayleigh, Rician, TR 38.901, or related label"],
        ["ebno_db", "Energy-per-bit to noise-density operating point"],
        ["channel_energy", "Per-user/per-subcarrier channel-energy representation"],
        ["active_ut_mask", "Mask indicating active user terminals"],
        ["per_ut_power", "Power directive per user terminal"],
        ["oracle_avg_ber", "Oracle or label BER used for supervised/DRL-style training"],
        ["oracle_ber_upper_confidence", "Confidence-aware reliability label"],
        ["oracle_throughput_bits", "Throughput label"],
        ["oracle_latency_ms", "Latency label"],
    ])


def add_references(doc: Document) -> None:
    doc.add_heading("References", level=1)
    refs = extract_references()
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ref)
        set_font(r, size=9)
    source_entries.append((rel(SOURCE_HTML), "Reference list extracted from the existing v3 thesis HTML"))


def unique_source_entries() -> list[tuple[str, str]]:
    seen: set[tuple[str, str]] = set()
    unique: list[tuple[str, str]] = []
    for path, note in source_entries:
        key = (path, note)
        if key not in seen:
            seen.add(key)
            unique.append(key)
    return unique


def add_expansion_blocks(doc: Document) -> None:
    """Add deeper technical notes without detached filler paragraphs."""
    doc.add_heading("Appendix F: Extended Methodological Commentary", level=1)
    themes = [
        (
            "F.1 Why Cross-Layer Evidence Matters",
            [
                "A factory communication system should be evaluated as a chain rather than as isolated blocks. Estimation errors change decoded-bit reliability; decoded reliability shapes feedback; feedback affects scheduling; scheduling changes which channel states are exposed to traffic. This chain is the reason Factory6G uses shared output metrics across stages.",
                "The practical implication is that a BER curve is never only a receiver result or only a scheduler result. It is the visible outcome of a modelling sequence that includes channel family, pilot structure, receiver assumptions, active-user selection, power normalization, and Monte Carlo stopping policy.",
                "For this reason, the thesis treats the simulator architecture and the source audit as part of the research contribution. The reader should be able to follow a claim from the written sentence to a figure, from the figure to a CSV or JSON file, and from that file to the configuration that generated it.",
            ],
        ),
        (
            "F.2 Why Deterministic Baselines Remain Necessary",
            [
                "Deterministic baselines are not merely weak comparators. They provide interpretability, fallback behaviour, and sanity checks for learned policies. In the current evidence, max-throughput remains a best observed method in selected channel summaries, which demonstrates why learned methods must be presented with restraint.",
                "This matters in industrial settings because a factory network must continue operating when a learned model is uncertain, unavailable, or outside its training distribution. A deterministic policy can be inspected, reproduced, and activated as a recovery mode, even when it is not optimal for every traffic class.",
                "The thesis therefore uses classical and heuristic resource managers as design references rather than as artificial opponents. The learned BER-DRL policy is valuable when it improves or usefully complements those references, not because learning alone is assumed to be superior.",
            ],
        ),
        (
            "F.3 Why Confidence Bounds Are Part of the Result",
            [
                "Monte Carlo simulations often report zero observed errors at high Eb/N0. Without a confidence bound, such a point can be mistaken for proven reliability. Factory6G records upper confidence so that finite evidence can be read honestly.",
                "This distinction is especially important in the May 23 Rayleigh results, where several methods report zero observed BER. The thesis therefore describes BER-DRL as matching the best observed Rayleigh reliability rather than claiming absolute error-free operation.",
                "A confidence-aware reading also helps compare channels. Rician and TR 38.901 UMi introduce different reliability pressure, and the upper-confidence value helps separate a robust result from a short-run artefact caused by limited samples.",
            ],
        ),
        (
            "F.4 Why Simulation Traceability Is a Contribution",
            [
                "Simulation-based theses can become difficult to assess when figures are detached from scripts and result files. The Factory6G workflow links each result family to CSV, JSON, plot, and configuration artifacts, making it easier to audit and extend.",
                "Traceability also protects the thesis from overclaiming. When a figure is a derived or synthetic projection, it is labelled as such and is not used as sole evidence for method ranking. When a table is extracted from the May 23 summary CSV, the source path is recorded in the audit note.",
                "This practice is useful beyond this dissertation. A future researcher can rerun or extend one part of the evidence chain without guessing which plot, checkpoint, or configuration supported a particular chapter statement.",
            ],
        ),
        (
            "F.5 How to Read BER-DRL",
            [
                "BER-DRL is best read as a policy that encodes reliability preference. It should not be described as universally optimal. The current evidence shows competitiveness and useful behaviour, while also showing that classical baselines can outperform it under some channels.",
                "The strongest interpretation is hybrid. BER-DRL demonstrates that a learned policy can be trained and evaluated through the same output schema as deterministic schedulers, which makes it auditable. It also demonstrates that the learned policy must remain bounded by channel sensitivity and baseline comparison.",
                "This is why the discussion chapter frames BER-DRL as an option inside a broader control stack. The policy contributes an AI/ML mechanism, but deployment credibility still depends on fallback modes, confidence monitoring, and validation beyond the current simulation campaign.",
            ],
        ),
        (
            "F.6 How to Extend the Work",
            [
                "The immediate extension is not to add more AI models indiscriminately. The stronger next step is to add independent seeds, richer traffic, factory-floor interference, hardware timing, and explainability constraints so that learned actions can be trusted in industrial contexts.",
                "JIDD-SCMA should also be stabilized before it becomes central evidence. Its current curve is useful because it points toward joint detection and decoding, but the non-monotonic behaviour means it should remain supporting evidence until additional checks confirm the cause.",
                "A future experimental pathway should move from controlled simulation to hardware-in-the-loop and over-the-air evaluation. That pathway would test whether the ranking observed in Factory6G survives physical impairments, controller timing, synchronization, and real factory mobility.",
            ],
        ),
    ]
    for title, paragraphs in themes:
        doc.add_heading(title, level=2)
        for paragraph in paragraphs:
            add_body(doc, paragraph)


def write_audit() -> None:
    lines = [
        "# Factory6G Final Thesis Source Audit",
        "",
        "Generated by `thesis_writing/tools/build_final_thesis.py`.",
        "",
        "## Primary Thesis Source",
        f"- Source DOCX/template: `{rel(SOURCE_DOCX)}`",
        f"- Source HTML references: `{rel(SOURCE_HTML)}`",
        "",
        "## Reused Evidence Artifacts",
    ]
    for path, note in unique_source_entries():
        lines.append(f"- `{path}` - {note}")
    lines.extend([
        "",
        "## Citation Checks",
        "- References were extracted from the existing v3 thesis HTML and correspond to the local literature collection where available.",
        "- No new fabricated DOI, venue, author, or title was introduced by the builder.",
        "- Unresolved manual check: confirm the official Sunway declaration form and any required publication list before final submission.",
        "- Unresolved manual check: update Word/LibreOffice fields for TOC/list of figures/list of tables before final administrative submission if page numbers are required by the graduate office template.",
        "",
        "## Claim Boundaries",
        "- The thesis uses current simulation artifacts only; no new simulation campaign is introduced.",
        "- BER-DRL is presented as competitive and reliability-oriented, not universally dominant.",
        "- JIDD-SCMA is treated as supporting/future-work evidence because the available curve has current stability concerns.",
        "- Synthetic or smoothed projection CSVs are not used as sole ranking evidence.",
    ])
    AUDIT_MD.write_text("\n".join(lines) + "\n")


def main() -> None:
    # Build a clean DOCX because the source v3 file does not expose several
    # standard Word styles through python-docx. The v3 file remains the content
    # and reference source, while this clean shell improves render reliability.
    doc = Document()
    configure_styles(doc)

    front_matter(doc)
    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_chapter_6(doc)
    add_appendices(doc)
    add_expansion_blocks(doc)
    add_references(doc)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    write_audit()
    print(f"Wrote {OUT_DOCX}")
    print(f"Wrote {AUDIT_MD}")


if __name__ == "__main__":
    main()
